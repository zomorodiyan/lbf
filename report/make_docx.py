#!/usr/bin/env python3
"""Build plc_effects_2page_may23.docx from plc_effects_2page_out.html.
Font size, line spacing, and margins match the PDF (9pt / 1.42 / 1.35–1.7 cm).
"""

import os, re
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/home/mzomoro1/bin/lbf3/report/plc_effects_2page_out.html"
OUT = "/home/mzomoro1/bin/lbf3/report/plc_effects_2page_may23.docx"

BODY   = 9.0   # pt — matches PDF body font
LS     = 1.42  # line spacing multiple — matches PDF

with open(SRC) as f:
    soup = BeautifulSoup(f.read(), "html.parser")

doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(1.35)
    sec.bottom_margin = Cm(1.35)
    sec.left_margin   = Cm(1.7)
    sec.right_margin  = Cm(1.7)

# ── Helpers ────────────────────────────────────────────────────────────────────

def sp(p, before=0, after=2):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.line_spacing = LS

def left_border(p, hex_color, shade_hex):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), hex_color)
    pBdr.append(left)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), shade_hex)
    pPr.append(shd)

def walk_inline(node, p, size=BODY, bold=False, italic=False):
    if isinstance(node, NavigableString):
        text = re.sub(r'\s+', ' ', str(node).replace("\xa0", " "))
        if text.strip():
            run = p.add_run(text)
            run.bold = bold; run.italic = italic
            run.font.size = Pt(size)
    elif isinstance(node, Tag):
        tag = node.name
        b = bold   or tag in ("strong", "b")
        i = italic or tag in ("em", "i")
        if tag == "sub":
            run = p.add_run(node.get_text())
            run.bold = b; run.italic = i
            run.font.size = Pt(size - 1.5)
            run.font.subscript = True
        elif tag == "sup":
            run = p.add_run(node.get_text())
            run.bold = b; run.italic = i
            run.font.size = Pt(size - 1.5)
            run.font.superscript = True
        else:
            for child in node.children:
                walk_inline(child, p, size, b, i)

def add_para(doc, el, size=BODY, after=2):
    p = doc.add_paragraph()
    sp(p, before=0, after=after)
    walk_inline(el, p, size)
    return p

def add_shaded(doc, el, shade_hex, border_color):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    sp(p, before=2, after=2)
    walk_inline(el, p, BODY)
    left_border(p, border_color, shade_hex)

def add_table(doc, tbl_el):
    rows = tbl_el.find_all("tr")
    if not rows:
        return
    ncols = max(len(r.find_all(["th","td"])) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row.find_all(["th","td"])):
            para = tbl.rows[ri].cells[ci].paragraphs[0]
            walk_inline(cell, para, size=BODY)
            for run in para.runs:
                if cell.name == "th":
                    run.bold = True

def add_figure(doc, fig_el):
    img_el = fig_el.find("img")
    cap_el = fig_el.find("figcaption")
    if img_el:
        path = os.path.join(os.path.dirname(SRC), img_el.get("src",""))
        if os.path.exists(path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sp(p, before=3, after=1)
            p.add_run().add_picture(path, width=Inches(5.2))
    if cap_el:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp(cp, before=0, after=4)
        walk_inline(cap_el, cp, size=BODY - 0.5)
        for run in cp.runs:
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def add_calc_block(doc, header, lines):
    """Bold label followed by indented equation/calculation lines."""
    p = doc.add_paragraph()
    sp(p, before=4, after=0)
    r = p.add_run(header)
    r.bold = True
    r.font.size = Pt(BODY)
    for line in lines:
        pl = doc.add_paragraph()
        pl.paragraph_format.left_indent = Cm(0.6)
        sp(pl, before=0, after=1)
        rl = pl.add_run(line)
        rl.font.size = Pt(BODY - 0.5)

def process(doc, el):
    if isinstance(el, NavigableString):
        return
    tag     = el.name
    classes = el.get("class", []) if isinstance(el, Tag) else []

    if tag == "h2":
        h = doc.add_heading(el.get_text().replace("\xa0"," "), level=2)
        for run in h.runs:
            run.font.size = Pt(10)
        h.paragraph_format.space_before = Pt(7)
        h.paragraph_format.space_after  = Pt(2)
        h.paragraph_format.line_spacing = LS

    elif tag == "h3":
        h = doc.add_heading(el.get_text().replace("\xa0"," "), level=3)
        for run in h.runs:
            run.font.size = Pt(9.5)
        h.paragraph_format.space_before = Pt(4)
        h.paragraph_format.space_after  = Pt(1)
        h.paragraph_format.line_spacing = LS

    elif tag == "p":
        add_para(doc, el)

    elif tag == "ul":
        for li in el.find_all("li", recursive=False):
            p = doc.add_paragraph(style="List Bullet")
            sp(p, before=0, after=1)
            walk_inline(li, p, size=BODY)

    elif tag == "ol":
        for i, li in enumerate(el.find_all("li", recursive=False), 1):
            p = doc.add_paragraph()
            sp(p, before=0, after=1)
            p.paragraph_format.left_indent       = Cm(0.6)
            p.paragraph_format.first_line_indent = Cm(-0.6)
            run = p.add_run(f"{i}. ")
            run.font.size = Pt(BODY)
            walk_inline(li, p, size=BODY)

    elif tag == "table":
        add_table(doc, el)

    elif tag == "figure":
        add_figure(doc, el)

    elif tag == "div":
        if "note" in classes:
            add_shaded(doc, el, "EEF4FF", "4A80C9")
        elif "assumption" in classes:
            add_shaded(doc, el, "F0F8EC", "4A904A")
        else:
            for child in el.children:
                process(doc, child)

# ── Equation/calculation injection functions ───────────────────────────────────
# µ=µ  ×=×  ≈=≈  ²=²  ³=³  ⁻=⁻
# ⁶=⁶  ⁸=⁸  ⁹=⁹  ¹=¹  ⁴=⁴  ⁰=⁰
# π=π  σ=σ  δ=δ  η=η  →=→  −=−  √=√

def inject_dissolution(doc):
    add_calc_block(doc, "Dissolution analysis:", [
        "Pore:     d = 50 µm  (R = 25 µm)",
        "Lifetime: τ_life ≈ 1 ms",
        "[%C] = 0.030,  [%O] = 0.001  (316L)",
    ])
    add_calc_block(doc, "H₂ — Sievert’s law:", [
        "Derive τ_H — the time it takes for H₂ gas inside a detached pore to fully dissolve into the surrounding liquid steel.",
        "[H]_wt% = K_H(T) √(p_H₂/p°);   K_H(1873 K) ≈ 28 cc/100 g  [4]",
        "C_eq ≈ 88 mol/m³  (from K_H, ρ_melt = 7000 kg/m³, p_H₂ = 1 atm)",
        "n₀ = p_H₂ V / RT = 101 kPa × (4/3)π(25 µm)³ / (8.314 × 1873) ≈ 4.3×10⁻¹³ mol",
        "δ = n₀ / (4πR² C_eq) ≈ 4.3×10⁻¹³ / (7.9×10⁻⁹ × 88) ≈ 0.62 µm",
        "τ_H = π δ² / D_H;   D_H ≈ 1.6×10⁻⁸ m²/s",
        "τ_H ≈ π × (0.62 µm)² / 1.6×10⁻⁸ ≈ 75 µs  <<  τ_life  →  dissolves completely",
    ])
    add_calc_block(doc, "CO — Chipman equilibrium:", [
        "Derive p_CO,eq — the equilibrium CO partial pressure the liquid steel can sustain given its dissolved carbon and oxygen content.",
        "log K_CO = 1160/T + 2.003;   K_CO = p_CO / ([%C][%O])",
        "At 1873 K:   p_CO,eq = K_CO × [%C][%O] = 419 × 0.030 × 0.001 ≈ 0.013 atm  (~80× below pore pressure)",
        "→  pore volume shrinks by ≈ (H₂+CO) fraction;  insoluble Ar persists as residual bubble",
    ])

def inject_morohoshi(doc):
    add_calc_block(doc, "Morohoshi (2013) Fe–C–O model:", [
        "K(T,w_C,w_O) = 10^(−0.17 w_O − 0.427 w_C) × w_O × exp(29300/T − 10.9)",
        "σ(T) = 1925 − 0.455(T − 1808) − 0.155 T ln(1 + K)  [mN/m]",
        "dσ/dT = −0.455 − 0.155 [ln(1+K) − K/(1+K) × 29300/T]  [mN/m·K]",
        "T* (dσ/dT = 0):   ≈1821 K  (η = 0.1%),   ≈1963 K  (η = 1%)",
    ])

def inject_energy(doc):
    add_calc_block(doc, "Energy calculation (PLC pyrolysis):", [
        "Strip width:  w = 1.5 × d_spot = 1.5 × 67 µm ≈ 100 µm",
        "E_L = P/v = 200 W / 200 mm·s⁻¹ = 1.0 J/mm",
        "V_PLC/mm = w × h_PLC × 1 mm = 100×50 µm² × 1 mm ≈ 5.0×10⁻¹² m³/mm",
        "m_PLC/mm = ρ_PLC × V = 1145 × 5.0×10⁻¹² ≈ 5.7×10⁻⁹ kg/mm",
        "H_d ≈ 10⁶ J/kg  (order-of-magnitude for polymer pyrolysis)",
        "E_PLC/E_L = m·H_d / E_L ≈ 5.7×10⁻³ / 1.0 ≈ 0.57%",
    ])

# ── Title + subtitle ───────────────────────────────────────────────────────────
h1 = soup.find("h1")
p  = doc.add_paragraph()
sp(p, before=0, after=2)
run = p.add_run(h1.get_text().replace("\xa0"," "))
run.bold = True; run.font.size = Pt(13)

sub = soup.find(class_="subtitle")
if sub:
    p2 = doc.add_paragraph(sub.get_text().replace("\xa0"," "))
    sp(p2, before=0, after=6)
    for r in p2.runs:
        r.font.size = Pt(BODY - 0.5)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# ── Main content — with selective paragraph merging and equation injection ──────
MERGE_TRIGGERS = [
    "Sievert",
    "At P\xa0=",
    "At P =",
]

children = [c for c in soup.find(class_="two-col").children
            if not (isinstance(c, NavigableString) and not c.strip())]

injected = set()

i = 0
while i < len(children):
    el = children[i]
    if (isinstance(el, Tag) and el.name == "p"
            and i + 1 < len(children)
            and isinstance(children[i+1], Tag) and children[i+1].name == "p"
            and any(t in el.get_text() for t in MERGE_TRIGGERS)):
        # Merge two consecutive <p>s into one Word paragraph
        p = doc.add_paragraph()
        sp(p, before=0, after=2)
        walk_inline(el, p, BODY)
        run = p.add_run("  ")
        run.font.size = Pt(BODY)
        walk_inline(children[i+1], p, BODY)

        # Inject dissolution equations after the merged H₂/CO para
        if "Sievert" in el.get_text() and "dissolution" not in injected:
            inject_dissolution(doc)
            injected.add("dissolution")

        # Inject energy calculation after the merged energy/absorptivity para
        if any(t in el.get_text() for t in ["At P\xa0=", "At P ="]) and "energy" not in injected:
            inject_energy(doc)
            injected.add("energy")

        i += 2
    else:
        process(doc, el)

        # Inject Morohoshi equations after the surface tension shift para
        if (isinstance(el, Tag) and el.name == "p"
                and "Dissolved C and O" in el.get_text()
                and "morohoshi" not in injected):
            inject_morohoshi(doc)
            injected.add("morohoshi")

        i += 1

# ── References ─────────────────────────────────────────────────────────────────
refs = soup.find(class_="refs-section")
if refs:
    process(doc, refs)

doc.save(OUT)
print(f"Saved: {OUT}")
